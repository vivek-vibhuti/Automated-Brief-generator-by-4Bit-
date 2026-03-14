import os
import PyPDF2
from docx import Document as DocxDocument
from typing import List, Dict, Optional
from ..models import Document, DocumentChunk
import pdfplumber  # Add this to requirements.txt
import re

class DocumentProcessor:
    """Enhanced document processor with better chunking and source tracking"""
    
    @staticmethod
    def process_pdf(file_path: str, document: Document) -> List[Dict]:
        """Enhanced PDF processing with pdfplumber for better text extraction"""
        chunks = []
        
        try:
            # Try pdfplumber first (better for complex PDFs)
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text() or ""
                    
                    if not text.strip():
                        # Try extracting with layout preservation
                        text = page.extract_text(layout=True) or ""
                    
                    if text.strip():
                        # Process this page's text into chunks
                        page_chunks = DocumentProcessor._chunk_text(
                            text=text,
                            source=document.title,
                            page=page_num,
                            doc_id=document.id
                        )
                        chunks.extend(page_chunks)
                    else:
                        print(f"⚠️ No text extracted from page {page_num} of {document.title}")
        
        except Exception as e:
            print(f"⚠️ pdfplumber failed, falling back to PyPDF2: {e}")
            # Fallback to PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text and text.strip():
                        page_chunks = DocumentProcessor._chunk_text(
                            text=text,
                            source=document.title,
                            page=page_num,
                            doc_id=document.id
                        )
                        chunks.extend(page_chunks)
        
        print(f"📄 PDF processed: {len(chunks)} chunks from {document.title}")
        return chunks
    
    @staticmethod
    def process_text(file_path: str, document: Document) -> List[Dict]:
        """Enhanced text file processing with smart chunking"""
        chunks = []
        
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            
            # Split into paragraphs first
            paragraphs = re.split(r'\n\s*\n', content)
            
            for i, para in enumerate(paragraphs):
                if para.strip():
                    # Further chunk long paragraphs
                    if len(para.split()) > 400:  # If paragraph is long
                        sub_chunks = DocumentProcessor._chunk_text(
                            text=para,
                            source=document.title,
                            page=None,
                            doc_id=document.id,
                            base_index=i
                        )
                        chunks.extend(sub_chunks)
                    else:
                        chunks.append({
                            'content': para.strip(),
                            'page_number': None,
                            'metadata': {
                                'document_id': document.id,
                                'document_title': document.title,
                                'chunk_index': i,
                                'paragraph': i,
                                'source': document.title,
                                'page': None
                            }
                        })
        
        print(f"📄 Text file processed: {len(chunks)} chunks from {document.title}")
        return chunks
    
    @staticmethod
    def process_docx(file_path: str, document: Document) -> List[Dict]:
        """Enhanced DOCX processing with paragraph grouping"""
        doc = DocxDocument(file_path)
        chunks = []
        
        # Group paragraphs into logical sections
        current_section = []
        current_section_text = ""
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                # Empty paragraph might indicate section break
                if current_section_text:
                    # Process the accumulated section
                    section_chunks = DocumentProcessor._chunk_text(
                        text=current_section_text,
                        source=document.title,
                        page=None,
                        doc_id=document.id,
                        base_index=len(chunks)
                    )
                    chunks.extend(section_chunks)
                    current_section = []
                    current_section_text = ""
                continue
            
            current_section.append(text)
            current_section_text += " " + text
            
            # Check if this might be a heading (short text, often with formatting)
            if len(text.split()) < 10 and paragraph.runs and paragraph.runs[0].bold:
                # This is likely a heading - treat as section boundary
                if current_section_text:
                    section_chunks = DocumentProcessor._chunk_text(
                        text=current_section_text,
                        source=document.title,
                        page=None,
                        doc_id=document.id,
                        base_index=len(chunks)
                    )
                    chunks.extend(section_chunks)
                    current_section = []
                    current_section_text = text  # Start new section with heading
        
        # Don't forget the last section
        if current_section_text:
            section_chunks = DocumentProcessor._chunk_text(
                text=current_section_text,
                source=document.title,
                page=None,
                doc_id=document.id,
                base_index=len(chunks)
            )
            chunks.extend(section_chunks)
        
        print(f"📄 DOCX processed: {len(chunks)} chunks from {document.title}")
        return chunks
    
    @staticmethod
    def _chunk_text(text: str, source: str, page: Optional[int], doc_id: int, 
                    base_index: int = 0, chunk_size: int = 500, overlap: int = 50) -> List[Dict]:
        """Intelligently chunk text with overlap for better context"""
        words = text.split()
        chunks = []
        
        if len(words) <= chunk_size:
            # Text is short enough to be one chunk
            chunks.append({
                'content': text.strip(),
                'page_number': page,
                'metadata': {
                    'document_id': doc_id,
                    'document_title': source,
                    'page': page,
                    'source': source,
                    'chunk_index': base_index,
                    'word_count': len(words)
                }
            })
        else:
            # Split into overlapping chunks
            for i in range(0, len(words), chunk_size - overlap):
                chunk_words = words[i:i + chunk_size]
                if len(chunk_words) < 50:  # Don't create tiny last chunks
                    # Merge with previous chunk
                    if chunks:
                        chunks[-1]['content'] += " " + " ".join(chunk_words)
                        chunks[-1]['metadata']['word_count'] += len(chunk_words)
                    break
                
                chunk_text = " ".join(chunk_words)
                chunks.append({
                    'content': chunk_text,
                    'page_number': page,
                    'metadata': {
                        'document_id': doc_id,
                        'document_title': source,
                        'page': page,
                        'source': source,
                        'chunk_index': base_index + i,
                        'word_count': len(chunk_words),
                        'chunk_start': i,
                        'chunk_end': i + len(chunk_words)
                    }
                })
        
        return chunks
    
    @classmethod
    def process_document(cls, document: Document) -> List[Dict]:
        """Process document based on file type with enhanced extraction"""
        file_path = document.file.path
        file_ext = os.path.splitext(file_path)[1].lower()
        
        print(f"🔄 Processing document: {document.title} ({file_ext})")
        
        try:
            if file_ext == '.pdf':
                chunks = cls.process_pdf(file_path, document)
            elif file_ext == '.txt':
                chunks = cls.process_text(file_path, document)
            elif file_ext == '.docx':
                chunks = cls.process_docx(file_path, document)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            # Add unique chunk IDs
            for i, chunk in enumerate(chunks):
                chunk['metadata']['chunk_id'] = f"{document.id}_{i}"
                chunk['metadata']['processed_at'] = str(document.uploaded_at)
            
            print(f"✅ Document processed: {len(chunks)} total chunks")
            return chunks
            
        except Exception as e:
            print(f"❌ Error processing document {document.title}: {e}")
            import traceback
            traceback.print_exc()
            raise
    
    @staticmethod
    def get_chunk_preview(chunk: Dict, max_length: int = 200) -> str:
        """Get a preview of chunk content for debugging"""
        content = chunk['content']
        if len(content) > max_length:
            return content[:max_length] + "..."
        return content


# Optional: Add a function to test the processor
def test_processor(file_path: str):
    """Test function to verify document processing"""
    from ..models import Document
    
    # Create a dummy document for testing
    class DummyDocument:
        def __init__(self, path, title):
            self.file = type('obj', (object,), {'path': path})()
            self.title = title
            self.id = 999
            self.uploaded_at = "2024-01-01"
    
    dummy_doc = DummyDocument(file_path, os.path.basename(file_path))
    processor = DocumentProcessor()
    chunks = processor.process_document(dummy_doc)
    
    print(f"\n📊 Processing Results:")
    print(f"   Total chunks: {len(chunks)}")
    if chunks:
        print(f"   First chunk preview: {processor.get_chunk_preview(chunks[0])}")
        print(f"   First chunk metadata: {chunks[0]['metadata']}")
    
    return chunks